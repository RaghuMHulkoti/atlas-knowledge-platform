"""
docx_parser.py

Microsoft Word (.docx) file parser backed by python-docx.

Extracts paragraph text and table cell text in document order. Legacy ``.doc``
(binary) files are not supported.
"""

import io

from docx import Document as DocxDocument

from app.core.exceptions import ConnectorException
from app.core.logging import get_logger
from app.infrastructure.connectors.files.base import BaseFileParser

logger = get_logger(__name__)


class DocxParser(BaseFileParser):
    """Extracts text from .docx files."""

    extensions = (".docx",)

    def extract_text(self, data: bytes, filename: str) -> str:
        try:
            document = DocxDocument(io.BytesIO(data))
        except Exception as exc:
            raise ConnectorException(
                f"Could not read DOCX '{filename}': {exc}"
            ) from exc

        parts: list[str] = [
            para.text for para in document.paragraphs if para.text.strip()
        ]

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        logger.info(
            "DocxParser: extracted %d text block(s) from '%s'.",
            len(parts),
            filename,
        )
        return "\n\n".join(parts)
